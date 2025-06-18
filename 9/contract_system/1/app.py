from flask import Flask, render_template, request, redirect, url_for, flash, send_file, jsonify
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
import os
import zipfile
import tempfile
import shutil
from docx import Document
from docx.shared import Inches
import re
from datetime import datetime

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///contract_management.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads/templates'
app.config['OUTPUT_FOLDER'] = 'outputs'

# Создаем папки если их нет
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)


# Модели базы данных
class Client(db.Model):
    __tablename__ = 'clients'

    id = db.Column(db.Integer, primary_key=True)
    number = db.Column(db.String(50), unique=True, nullable=False)
    full_name = db.Column(db.String(500), nullable=False)
    short_name = db.Column(db.String(200))
    inn = db.Column(db.String(12), nullable=False)
    ogrn = db.Column(db.String(15))
    address = db.Column(db.Text)
    position = db.Column(db.String(200))
    position_genitive = db.Column(db.String(200))
    representative_name = db.Column(db.String(200))
    representative_name_genitive = db.Column(db.String(200))
    basis = db.Column(db.String(500))
    bank_details = db.Column(db.Text)
    phone = db.Column(db.String(50))
    email = db.Column(db.String(100))
    website = db.Column(db.String(200))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class Organization(db.Model):
    __tablename__ = 'organizations'

    id = db.Column(db.Integer, primary_key=True)
    signatory_position = db.Column(db.String(200), nullable=False)
    signatory_name = db.Column(db.String(200), nullable=False)
    signatory_power_of_attorney = db.Column(db.String(500))
    executor_position = db.Column(db.String(200), nullable=False)
    executor_name = db.Column(db.String(200), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class ContractTemplate(db.Model):
    __tablename__ = 'contract_templates'

    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    # Связь с файлами шаблонов
    files = db.relationship('TemplateFile', backref='template', lazy=True, cascade='all, delete-orphan')


class TemplateFile(db.Model):
    __tablename__ = 'template_files'

    id = db.Column(db.Integer, primary_key=True)
    template_id = db.Column(db.Integer, db.ForeignKey('contract_templates.id'), nullable=False)
    filename = db.Column(db.String(200), nullable=False)
    original_filename = db.Column(db.String(200), nullable=False)
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow)


# Создание таблиц
with app.app_context():
    db.create_all()


# Вспомогательные функции
def replace_placeholders(text, data):
    """Заменяет плейсхолдеры в тексте на данные"""
    if not text:
        return text

    for key, value in data.items():
        placeholder = '{{ ' + key + ' }}'
        if placeholder in text:
            text = text.replace(placeholder, str(value or ''))
    return text


def normalize_url(url):
    """Нормализует URL, добавляя протокол если необходимо"""
    if not url:
        return url

    url = url.strip()
    if url and not url.startswith(('http://', 'https://')):
        return f'https://{url}'
    return url


def clean_filename(filename):
    """Очищает имя файла от недопустимых символов, сохраняя читаемость"""
    # Убираем расширение
    name_without_ext = os.path.splitext(filename)[0] if '.' in filename else filename
    # Заменяем недопустимые символы на подчеркивания
    clean_name = re.sub(r'[<>:"/\\|?*]', '_', name_without_ext)
    # Убираем лишние пробелы и подчеркивания
    clean_name = re.sub(r'[_\s]+', '_', clean_name).strip('_')
    return clean_name if clean_name else 'document'


def simple_replace_in_paragraph(paragraph, data):
    """Безопасная замена плейсхолдеров в параграфе с улучшенной обработкой ошибок"""
    if not paragraph or not hasattr(paragraph, 'runs'):
        return False

    # Проверяем наличие runs
    if not paragraph.runs:
        return False

    # Получаем весь текст параграфа
    full_text = paragraph.text
    if not full_text or '{{' not in full_text:
        return False

    # Заменяем плейсхолдеры
    original_text = full_text
    for key, value in data.items():
        placeholder = '{{ ' + key + ' }}'
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, str(value or ''))

    # Если изменений не было
    if full_text == original_text:
        return False

    # Безопасная замена: проверяем наличие runs
    try:
        if len(paragraph.runs) > 0:
            # Сохраняем форматирование первого run'а
            first_run = paragraph.runs[0]
            first_run.text = full_text

            # Очищаем остальные runs (в обратном порядке, чтобы не сбить индексы)
            for i in range(len(paragraph.runs) - 1, 0, -1):
                try:
                    paragraph.runs[i].text = ""
                except IndexError:
                    # Пропускаем, если индекс уже недоступен
                    continue

            return True
        else:
            # Если нет runs, создаем новый
            run = paragraph.add_run(full_text)
            return True

    except Exception as e:
        print(f"Ошибка при замене в параграфе: {e}")
        # Попробуем создать новый run с текстом
        try:
            # Очищаем все существующие runs
            for run in paragraph.runs:
                run.text = ""
            # Добавляем новый run с замененным текстом
            paragraph.add_run(full_text)
            return True
        except Exception as e2:
            print(f"Критическая ошибка при создании run: {e2}")
            return False


def process_docx_template_safe(template_path, output_path, data):
    """Безопасная версия обработки DOCX с улучшенной обработкой ошибок"""
    try:
        print(f"📄 Открытие файла: {template_path}")
        doc = Document(template_path)

        replacements_made = 0

        # Обработка основного текста
        print("🔍 Поиск плейсхолдеров в основном тексте...")
        for i, paragraph in enumerate(doc.paragraphs):
            try:
                if simple_replace_in_paragraph(paragraph, data):
                    replacements_made += 1
                    print(f"  ✅ Замена в параграфе {i}")
            except Exception as e:
                print(f"  ⚠️ Ошибка в параграфе {i}: {e}")
                continue

        # Обработка таблиц
        print("🔍 Поиск плейсхолдеров в таблицах...")
        for table_idx, table in enumerate(doc.tables):
            try:
                for row_idx, row in enumerate(table.rows):
                    try:
                        for cell_idx, cell in enumerate(row.cells):
                            try:
                                for para_idx, paragraph in enumerate(cell.paragraphs):
                                    try:
                                        if simple_replace_in_paragraph(paragraph, data):
                                            replacements_made += 1
                                            print(
                                                f"  ✅ Замена в таблице {table_idx}, строка {row_idx}, ячейка {cell_idx}")
                                    except Exception as e:
                                        print(
                                            f"  ⚠️ Ошибка в параграфе таблицы {table_idx}[{row_idx}][{cell_idx}][{para_idx}]: {e}")
                                        continue
                            except Exception as e:
                                print(f"  ⚠️ Ошибка в ячейке таблицы {table_idx}[{row_idx}][{cell_idx}]: {e}")
                                continue
                    except Exception as e:
                        print(f"  ⚠️ Ошибка в строке таблицы {table_idx}[{row_idx}]: {e}")
                        continue
            except Exception as e:
                print(f"  ⚠️ Ошибка в таблице {table_idx}: {e}")
                continue

        # Обработка заголовков и футеров
        print("🔍 Поиск плейсхолдеров в заголовках и футерах...")
        for section_idx, section in enumerate(doc.sections):
            try:
                # Заголовки
                if hasattr(section, 'header') and section.header:
                    for para_idx, paragraph in enumerate(section.header.paragraphs):
                        try:
                            if simple_replace_in_paragraph(paragraph, data):
                                replacements_made += 1
                                print(f"  ✅ Замена в заголовке секции {section_idx}")
                        except Exception as e:
                            print(f"  ⚠️ Ошибка в заголовке секции {section_idx}[{para_idx}]: {e}")
                            continue

                # Футеры
                if hasattr(section, 'footer') and section.footer:
                    for para_idx, paragraph in enumerate(section.footer.paragraphs):
                        try:
                            if simple_replace_in_paragraph(paragraph, data):
                                replacements_made += 1
                                print(f"  ✅ Замена в футере секции {section_idx}")
                        except Exception as e:
                            print(f"  ⚠️ Ошибка в футере секции {section_idx}[{para_idx}]: {e}")
                            continue
            except Exception as e:
                print(f"  ⚠️ Ошибка в секции {section_idx}: {e}")
                continue

        print(f"📊 Всего замен выполнено: {replacements_made}")

        if replacements_made == 0:
            print("⚠️  ВНИМАНИЕ: Ни одного плейсхолдера не найдено!")
            print("Проверьте формат плейсхолдеров в документе. Должно быть: {{ Имя_переменной }}")

            # Показываем весь текст документа для отладки
            print("\n📋 Содержимое документа для отладки (первые 10 параграфов):")
            for i, para in enumerate(doc.paragraphs[:10]):
                if para.text.strip():
                    print(f"  Параграф {i}: {para.text}")

        # Сохраняем документ
        doc.save(output_path)
        print(f"💾 Документ сохранен: {output_path}")

        return True

    except Exception as e:
        print(f"❌ Критическая ошибка при обработке документа: {e}")
        import traceback
        traceback.print_exc()
        raise e


def diagnose_document_placeholders(template_path):
    """Диагностика плейсхолдеров в документе"""
    try:
        doc = Document(template_path)
        placeholders_found = set()

        # Поиск в основном тексте
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Ищем все плейсхолдеры в формате {{ text }}
            matches = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
            for match in matches:
                placeholders_found.add(match.strip())

        # Поиск в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        matches = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
                        for match in matches:
                            placeholders_found.add(match.strip())

        # Поиск в заголовках и футерах
        for section in doc.sections:
            if hasattr(section, 'header') and section.header:
                for paragraph in section.header.paragraphs:
                    text = paragraph.text
                    matches = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
                    for match in matches:
                        placeholders_found.add(match.strip())

            if hasattr(section, 'footer') and section.footer:
                for paragraph in section.footer.paragraphs:
                    text = paragraph.text
                    matches = re.findall(r'\{\{\s*([^}]+)\s*\}\}', text)
                    for match in matches:
                        placeholders_found.add(match.strip())

        return list(placeholders_found)

    except Exception as e:
        print(f"Ошибка диагностики: {e}")
        return []


def calculate_file_hash(file_path):
    """Вычисляет хэш файла для проверки уникальности"""
    import hashlib
    hash_md5 = hashlib.md5()
    try:
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    except Exception as e:
        print(f"Ошибка при вычислении хэша файла {file_path}: {e}")
        return "error"


def debug_file_contents(file_path, max_chars=500):
    """Показывает начало содержимого файла для отладки"""
    try:
        doc = Document(file_path)
        text_content = []

        # Собираем текст из первых нескольких параграфов
        for i, paragraph in enumerate(doc.paragraphs[:5]):
            if paragraph.text.strip():
                text_content.append(f"П{i}: {paragraph.text[:100]}...")

        content_preview = " | ".join(text_content)
        return content_preview[:max_chars] if content_preview else "Нет текстового содержимого"
    except Exception as e:
        return f"Ошибка чтения: {e}"


def debug_template_processing(templates, template_ids):
    """Отладочная функция для анализа выбранных шаблонов"""
    print(f"\n🔍 ОТЛАДКА ШАБЛОНОВ:")
    print(f"   Получено ID шаблонов: {template_ids}")
    print(f"   Найдено шаблонов в БД: {len(templates)}")

    for i, template in enumerate(templates):
        print(f"\n   📋 Шаблон #{i + 1}:")
        print(f"      ID: {template.id}")
        print(f"      Название: '{template.name}'")
        print(f"      Описание: '{template.description or 'Не указано'}'")
        print(f"      Дата создания: {template.created_at}")
        print(f"      Количество файлов: {len(template.files)}")

        for j, template_file in enumerate(template.files):
            file_path = os.path.join('uploads/templates', template_file.filename)
            file_exists = os.path.exists(file_path)
            file_size = os.path.getsize(file_path) if file_exists else 0
            file_hash = calculate_file_hash(file_path) if file_exists else "N/A"

            print(f"         📄 Файл #{j + 1}:")
            print(f"            ID в БД: {template_file.id}")
            print(f"            Оригинал: {template_file.original_filename}")
            print(f"            Система: {template_file.filename}")
            print(f"            Существует: {'✅' if file_exists else '❌'}")
            print(f"            Размер: {file_size} байт")
            print(f"            MD5: {file_hash}")
            print(f"            Загружен: {template_file.uploaded_at}")

            # Проверяем уникальность файлов
            if file_exists:
                try:
                    content_preview = debug_file_contents(file_path, 100)
                    print(f"            Содержимое: {content_preview}")
                except Exception as e:
                    print(f"            Ошибка чтения: {e}")

        # Проверяем, есть ли дублирующиеся хэши в одном шаблоне
        file_hashes = []
        for template_file in template.files:
            file_path = os.path.join('uploads/templates', template_file.filename)
            if os.path.exists(file_path):
                file_hash = calculate_file_hash(file_path)
                file_hashes.append((template_file.original_filename, file_hash))

        # Проверяем дубликаты
        seen_hashes = set()
        duplicates = []
        for filename, hash_val in file_hashes:
            if hash_val in seen_hashes:
                duplicates.append(filename)
            seen_hashes.add(hash_val)

        if duplicates:
            print(f"      ⚠️ НАЙДЕНЫ ДУБЛИРУЮЩИЕСЯ ФАЙЛЫ: {duplicates}")
        else:
            print(f"      ✅ Все файлы уникальны")


@app.route('/')
def index():
    return render_template('index.html')


# Маршруты для клиентов
@app.route('/clients')
def clients():
    clients = Client.query.all()
    return render_template('clients.html', clients=clients)


@app.route('/clients/add', methods=['GET', 'POST'])
def add_client():
    if request.method == 'POST':
        client = Client(
            number=request.form['number'],
            full_name=request.form['full_name'],
            short_name=request.form['short_name'],
            inn=request.form['inn'],
            ogrn=request.form['ogrn'],
            address=request.form['address'],
            position=request.form['position'],
            position_genitive=request.form['position_genitive'],
            representative_name=request.form['representative_name'],
            representative_name_genitive=request.form['representative_name_genitive'],
            basis=request.form['basis'],
            bank_details=request.form['bank_details'],
            phone=request.form['phone'],
            email=request.form['email'],
            website=normalize_url(request.form['website'])
        )

        try:
            db.session.add(client)
            db.session.commit()
            flash('Клиент успешно добавлен!', 'success')
            return redirect(url_for('clients'))
        except Exception as e:
            db.session.rollback()
            flash(f'Ошибка при добавлении клиента: {str(e)}', 'error')

    return render_template('add_client.html')


@app.route('/clients/edit/<int:id>', methods=['GET', 'POST'])
def edit_client(id):
    client = Client.query.get_or_404(id)

    if request.method == 'POST':
        client.number = request.form['number']
        client.full_name = request.form['full_name']
        client.short_name = request.form['short_name']
        client.inn = request.form['inn']
        client.ogrn = request.form['ogrn']
        client.address = request.form['address']
        client.position = request.form['position']
        client.position_genitive = request.form['position_genitive']
        client.representative_name = request.form['representative_name']
        client.representative_name_genitive = request.form['representative_name_genitive']
        client.basis = request.form['basis']
        client.bank_details = request.form['bank_details']
        client.phone = request.form['phone']
        client.email = request.form['email']
        client.website = normalize_url(request.form['website'])

        try:
            db.session.commit()
            flash('Клиент успешно обновлен!', 'success')
            return redirect(url_for('clients'))
        except Exception as e:
            db.session.rollback()
            flash(f'Ошибка при обновлении клиента: {str(e)}', 'error')

    return render_template('edit_client.html', client=client)


@app.route('/clients/delete/<int:id>')
def delete_client(id):
    client = Client.query.get_or_404(id)
    try:
        db.session.delete(client)
        db.session.commit()
        flash('Клиент успешно удален!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Ошибка при удалении клиента: {str(e)}', 'error')

    return redirect(url_for('clients'))


# Маршруты для организаций
@app.route('/organizations')
def organizations():
    organizations = Organization.query.all()
    return render_template('organizations.html', organizations=organizations)


@app.route('/organizations/add', methods=['GET', 'POST'])
def add_organization():
    if request.method == 'POST':
        organization = Organization(
            signatory_position=request.form['signatory_position'],
            signatory_name=request.form['signatory_name'],
            signatory_power_of_attorney=request.form['signatory_power_of_attorney'],
            executor_position=request.form['executor_position'],
            executor_name=request.form['executor_name']
        )

        try:
            db.session.add(organization)
            db.session.commit()
            flash('Организация успешно добавлена!', 'success')
            return redirect(url_for('organizations'))
        except Exception as e:
            db.session.rollback()
            flash(f'Ошибка при добавлении организации: {str(e)}', 'error')

    return render_template('add_organization.html')


@app.route('/organizations/edit/<int:id>', methods=['GET', 'POST'])
def edit_organization(id):
    organization = Organization.query.get_or_404(id)

    if request.method == 'POST':
        organization.signatory_position = request.form['signatory_position']
        organization.signatory_name = request.form['signatory_name']
        organization.signatory_power_of_attorney = request.form['signatory_power_of_attorney']
        organization.executor_position = request.form['executor_position']
        organization.executor_name = request.form['executor_name']

        try:
            db.session.commit()
            flash('Организация успешно обновлена!', 'success')
            return redirect(url_for('organizations'))
        except Exception as e:
            db.session.rollback()
            flash(f'Ошибка при обновлении организации: {str(e)}', 'error')

    return render_template('edit_organization.html', organization=organization)


@app.route('/organizations/delete/<int:id>')
def delete_organization(id):
    organization = Organization.query.get_or_404(id)
    try:
        db.session.delete(organization)
        db.session.commit()
        flash('Организация успешно удалена!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Ошибка при удалении организации: {str(e)}', 'error')

    return redirect(url_for('organizations'))


# Маршруты для шаблонов
@app.route('/templates')
def templates():
    templates = ContractTemplate.query.all()
    return render_template('templates.html', templates=templates)


@app.route('/templates/upload', methods=['POST'])
def upload_template():
    if 'files' not in request.files:
        flash('Файлы не выбраны', 'error')
        return redirect(url_for('templates'))

    files = request.files.getlist('files')
    template_name = request.form.get('name', '').strip()

    if not template_name:
        flash('Укажите название шаблона', 'error')
        return redirect(url_for('templates'))

    valid_files = []
    for file in files:
        if file.filename != '' and file.filename.endswith('.docx'):
            valid_files.append(file)

    if not valid_files:
        flash('Выберите хотя бы один файл .docx', 'error')
        return redirect(url_for('templates'))

    try:
        # Создаем новый шаблон
        template = ContractTemplate(
            name=template_name,
            description=request.form.get('description', '')
        )
        db.session.add(template)
        db.session.flush()  # Получаем ID шаблона

        # Сохраняем файлы
        import time
        import random

        for i, file in enumerate(valid_files):
            filename = secure_filename(file.filename)
            # Добавляем timestamp с микросекундами и случайное число для гарантии уникальности
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
            random_suffix = random.randint(1000, 9999)
            unique_filename = f"{timestamp}_{random_suffix}_{i:02d}_{filename}"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)

            print(f"📁 Сохранение файла #{i + 1}:")
            print(f"   Оригинальное имя: {filename}")
            print(f"   Уникальное имя: {unique_filename}")
            print(f"   Путь сохранения: {filepath}")

            file.save(filepath)

            # Проверяем, что файл действительно сохранился
            if os.path.exists(filepath):
                saved_size = os.path.getsize(filepath)
                saved_hash = calculate_file_hash(filepath)
                print(f"   ✅ Файл сохранен, размер: {saved_size} байт")
                print(f"   🔐 MD5: {saved_hash}")
            else:
                print(f"   ❌ ОШИБКА: Файл не сохранился!")

            # Создаем запись о файле
            template_file = TemplateFile(
                template_id=template.id,
                filename=unique_filename,
                original_filename=filename
            )
            db.session.add(template_file)
            print(f"   📝 Запись в БД создана: ID шаблона {template.id}, файл {unique_filename}")

            # Небольшая задержка для гарантии разных timestamp'ов
            time.sleep(0.01)

        db.session.commit()
        flash(f'Шаблон "{template_name}" успешно загружен с {len(valid_files)} файлами!', 'success')

    except Exception as e:
        db.session.rollback()
        flash(f'Ошибка при сохранении шаблона: {str(e)}', 'error')

    return redirect(url_for('templates'))


@app.route('/templates/delete/<int:id>')
def delete_template(id):
    template = ContractTemplate.query.get_or_404(id)

    try:
        # Удаляем файлы с диска
        for template_file in template.files:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], template_file.filename)
            if os.path.exists(file_path):
                os.remove(file_path)

        # Удаляем из базы данных (файлы удалятся автоматически из-за cascade)
        db.session.delete(template)
        db.session.commit()
        flash('Шаблон успешно удален!', 'success')

    except Exception as e:
        db.session.rollback()
        flash(f'Ошибка при удалении шаблона: {str(e)}', 'error')

    return redirect(url_for('templates'))


@app.route('/templates/test/<int:id>')
def test_template(id):
    """Тестирование шаблона без генерации договора"""
    template = ContractTemplate.query.get_or_404(id)

    print(f"\n=== ТЕСТ ШАБЛОНА '{template.name}' ===")

    result = {
        'template_name': template.name,
        'template_description': template.description,
        'files_count': len(template.files),
        'files': []
    }

    for template_file in template.files:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], template_file.filename)

        file_info = {
            'original_filename': template_file.original_filename,
            'system_filename': template_file.filename,
            'exists': os.path.exists(file_path),
            'size': 0,
            'hash': '',
            'placeholders': [],
            'content_preview': '',
            'error': None
        }

        if os.path.exists(file_path):
            try:
                file_info['size'] = os.path.getsize(file_path)
                file_info['hash'] = calculate_file_hash(file_path)
                file_info['placeholders'] = diagnose_document_placeholders(file_path)
                file_info['content_preview'] = debug_file_contents(file_path)

                print(f"\n📄 Файл: {template_file.original_filename}")
                print(f"   Размер: {file_info['size']} байт")
                print(f"   MD5: {file_info['hash']}")
                print(f"   Плейсхолдеры: {len(file_info['placeholders'])}")
                print(f"   Содержимое: {file_info['content_preview'][:100]}...")

            except Exception as e:
                file_info['error'] = str(e)
                print(f"   ❌ Ошибка анализа: {e}")
        else:
            file_info['error'] = 'Файл не найден'
            print(f"   ❌ Файл не найден: {file_path}")

        result['files'].append(file_info)

    return jsonify(result)


@app.route('/templates/debug/<int:id>')
def debug_template_files(id):
    """Отладочная функция для детального анализа файлов в шаблоне"""
    template = ContractTemplate.query.get_or_404(id)

    result = {
        'template': {
            'id': template.id,
            'name': template.name,
            'description': template.description,
            'files_count': len(template.files)
        },
        'files': [],
        'analysis': {
            'duplicate_hashes': []
        }
    }

    file_hashes = {}

    for template_file in template.files:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], template_file.filename)

        file_info = {
            'id': template_file.id,
            'original_filename': template_file.original_filename,
            'system_filename': template_file.filename,
            'uploaded_at': template_file.uploaded_at.isoformat(),
            'exists': os.path.exists(file_path),
            'size': 0,
            'hash': None,
            'content_preview': None,
            'placeholders': []
        }

        if os.path.exists(file_path):
            try:
                file_info['size'] = os.path.getsize(file_path)
                file_info['hash'] = calculate_file_hash(file_path)
                file_info['content_preview'] = debug_file_contents(file_path, 300)
                file_info['placeholders'] = diagnose_document_placeholders(file_path)

                # Проверяем дубликаты
                if file_info['hash'] in file_hashes:
                    result['analysis']['duplicate_hashes'].append({
                        'hash': file_info['hash'],
                        'files': [file_hashes[file_info['hash']], template_file.original_filename]
                    })
                else:
                    file_hashes[file_info['hash']] = template_file.original_filename

            except Exception as e:
                file_info['error'] = str(e)

        result['files'].append(file_info)

    return jsonify(result)


@app.route('/templates/diagnose/<int:id>')
def diagnose_template(id):
    """Диагностика шаблона"""
    template = ContractTemplate.query.get_or_404(id)

    result = {
        'template_name': template.name,
        'files': []
    }

    for template_file in template.files:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], template_file.filename)
        if os.path.exists(file_path):
            placeholders = diagnose_document_placeholders(file_path)
            result['files'].append({
                'filename': template_file.original_filename,
                'placeholders': placeholders
            })
        else:
            result['files'].append({
                'filename': template_file.original_filename,
                'error': 'Файл не найден'
            })

    return jsonify(result)


# Маршруты для генерации договоров
@app.route('/contracts')
def contracts():
    clients = Client.query.all()
    organizations = Organization.query.all()
    templates = ContractTemplate.query.all()
    return render_template('contracts.html', clients=clients, organizations=organizations, templates=templates)


@app.route('/contracts/generate', methods=['POST'])
def generate_contracts():
    client_id = request.form['client_id']
    organization_id = request.form['organization_id']
    template_ids = request.form.getlist('template_ids')

    if not client_id or not organization_id or not template_ids:
        flash('Необходимо выбрать клиента, организацию и шаблоны', 'error')
        return redirect(url_for('contracts'))

    client = Client.query.get_or_404(client_id)
    organization = Organization.query.get_or_404(organization_id)
    templates = ContractTemplate.query.filter(ContractTemplate.id.in_(template_ids)).all()

    # Отладочная информация о шаблонах
    debug_template_processing(templates, template_ids)

    # Подготовка данных для замены
    data = {
        'Номер': client.number,
        'Полное_наименование': client.full_name,
        'Сокращенное_наименование': client.short_name,
        'ИНН': client.inn,
        'ОГРН': client.ogrn,
        'Адрес': client.address,
        'Должность': client.position,
        'Должность_р': client.position_genitive,
        'ФИО_представителя': client.representative_name,
        'ФИО_представителя_р': client.representative_name_genitive,
        'Основание': client.basis,
        'Реквизиты': client.bank_details,
        'Телефон': client.phone,
        'Электронная_почта': client.email,
        'Сайт': client.website,
        'Должность_подписанта': organization.signatory_position,
        'ФИО_подписанта': organization.signatory_name,
        'Доверенность_подписанта': organization.signatory_power_of_attorney,
        'Должность_исполнителя': organization.executor_position,
        'ФИО_исполнителя': organization.executor_name,
        'Дата': datetime.now().strftime('%d.%m.%Y')
    }

    # Отладочная информация
    print("\n=== ГЕНЕРАЦИЯ ДОГОВОРОВ ===")
    print(f"Клиент: {client.full_name}")
    print(f"Организация: {organization.signatory_name}")
    print(f"Шаблонов выбрано: {len(templates)}")
    print(f"ID шаблонов: {template_ids}")
    print("\nДанные для замены:")
    for key, value in data.items():
        print(f"  {{ {key} }} = '{value}'")
    print()

    # Создание временной папки для файлов
    with tempfile.TemporaryDirectory() as temp_dir:
        generated_files = []
        used_filenames = set()  # Множество для отслеживания уже использованных имен
        base_timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

        # Подготавливаем очищенное имя клиента для архива
        client_short = client.short_name or client.full_name
        clean_client_name = clean_filename(client_short)

        for template_idx, template in enumerate(templates):
            print(f"\n🔄 Обработка шаблона #{template_idx + 1}: '{template.name}' (ID: {template.id})")
            print(f"   Описание: {template.description or 'Не указано'}")
            print(f"   Файлов в шаблоне: {len(template.files)}")

            if not template.files:
                print(f"   ⚠️ ВНИМАНИЕ: В шаблоне '{template.name}' нет файлов!")
                continue

            for file_idx, template_file in enumerate(template.files):
                template_path = os.path.join(app.config['UPLOAD_FOLDER'], template_file.filename)

                print(f"\n   📁 Файл #{file_idx + 1} в шаблоне:")
                print(f"      ID файла в БД: {template_file.id}")
                print(f"      Оригинальное имя: {template_file.original_filename}")
                print(f"      Системное имя: {template_file.filename}")
                print(f"      Путь к файлу: {template_path}")
                print(f"      Дата загрузки: {template_file.uploaded_at}")

                if not os.path.exists(template_path):
                    print(f"      ❌ ОШИБКА: Файл не найден!")
                    flash(f'Файл шаблона не найден: {template_file.original_filename}', 'error')
                    continue

                # Дополнительная проверка файла
                file_size = os.path.getsize(template_path)
                file_hash = calculate_file_hash(template_path)
                print(f"      📊 Размер исходного файла: {file_size} байт")
                print(f"      🔐 MD5 исходного файла: {file_hash}")

                # Проверяем содержимое файла
                try:
                    doc_preview = debug_file_contents(template_path, 200)
                    print(f"      📄 Предварительный просмотр: {doc_preview}")
                except Exception as e:
                    print(f"      ⚠️ Ошибка чтения содержимого: {e}")

                # Формируем уникальное имя файла
                clean_original_name = clean_filename(template_file.original_filename)

                # Базовое имя файла
                base_filename = f"{clean_client_name}_{clean_original_name}_{base_timestamp}"
                output_filename = f"{base_filename}.docx"

                # Проверяем уникальность и добавляем номер если нужно
                counter = 1
                while output_filename in used_filenames:
                    output_filename = f"{base_filename}_{counter:02d}.docx"
                    counter += 1

                # Добавляем имя в множество использованных
                used_filenames.add(output_filename)
                output_path = os.path.join(temp_dir, output_filename)

                print(f"      📤 Выходной файл: {output_filename}")

                try:
                    print(f"      🔄 Начало обработки...")

                    # Используем безопасную версию обработки
                    process_docx_template_safe(template_path, output_path, data)

                    # Проверяем, что файл действительно создан
                    if os.path.exists(output_path):
                        file_size = os.path.getsize(output_path)
                        output_hash = calculate_file_hash(output_path)
                        generated_files.append(output_path)
                        print(f"      ✅ Файл успешно создан! Размер: {file_size} байт")
                        print(f"      🔐 MD5 выходного файла: {output_hash}")

                        # Сравниваем хэши
                        if output_hash == file_hash:
                            print(f"      ⚠️ ВНИМАНИЕ: Выходной файл идентичен исходному (замены не произошло)")
                        else:
                            print(f"      ✅ Файл успешно обработан (хэш изменился)")

                        # Проверяем содержимое выходного файла
                        try:
                            output_preview = debug_file_contents(output_path, 200)
                            print(f"      📄 Содержимое выходного файла: {output_preview}")
                        except Exception as e:
                            print(f"      ⚠️ Ошибка чтения выходного файла: {e}")
                    else:
                        print(f"      ❌ Файл не был создан!")
                        flash(f'Не удалось создать файл для {template_file.original_filename}', 'error')

                except Exception as e:
                    print(f"      💥 ОШИБКА при обработке: {str(e)}")
                    print("      📋 Детали ошибки:")
                    import traceback
                    traceback.print_exc()
                    flash(
                        f'Ошибка при обработке файла {template_file.original_filename} из шаблона {template.name}: {str(e)}',
                        'error')
                    continue

        print(f"\n📊 ИТОГИ ГЕНЕРАЦИИ:")
        print(f"   Всего создано файлов: {len(generated_files)}")

        if generated_files:
            print("   📋 Список созданных файлов:")
            for i, file_path in enumerate(generated_files, 1):
                file_name = os.path.basename(file_path)
                file_size = os.path.getsize(file_path)
                print(f"      {i}. {file_name} ({file_size} байт)")

            # Создание архива
            archive_name = f"contracts_{clean_client_name}_{base_timestamp}.zip"
            archive_path = os.path.join(app.config['OUTPUT_FOLDER'], archive_name)

            print(f"\n📦 Создание архива: {archive_name}")
            try:
                with zipfile.ZipFile(archive_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for file_path in generated_files:
                        file_name = os.path.basename(file_path)
                        zipf.write(file_path, file_name)
                        print(f"      ✅ Добавлен в архив: {file_name}")

                archive_size = os.path.getsize(archive_path)
                print(f"📦 Архив создан успешно! Размер: {archive_size} байт")
                return send_file(archive_path, as_attachment=True, download_name=archive_name)

            except Exception as e:
                print(f"❌ Ошибка при создании архива: {e}")
                flash(f'Ошибка при создании архива: {str(e)}', 'error')
                return redirect(url_for('contracts'))
        else:
            print("❌ Не удалось создать ни одного файла")
            flash('Не удалось создать ни одного договора. Проверьте шаблоны и данные.', 'error')
            return redirect(url_for('contracts'))


if __name__ == '__main__':
    app.run(debug=True)